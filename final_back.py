from fastapi import FastAPI, Request, File, UploadFile, Form
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from io import BytesIO
import smtplib
from email.message import EmailMessage
import requests
from typing import List
import base64                   # uvicorn final_back:app --reload
import mimetypes

app = FastAPI()

# Allow CORS (so HTML/JS can call these APIs)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Replace with real values
SENDER_EMAIL = "email"
SENDER_APP_PASSWORD = "password"
GROQ_API_KEY = "api-key"
MODEL_NAME = "llama-3.3-70b-versatile"

class ReportData(BaseModel):
    language: str   # Default to English
    date: str
    time: str
    location: str
    person: str
    reporter: str
    recipient: str
    description: str
    evidence: str = ""

class EmailData(BaseModel):
    report: str
    emails: str

@app.post("/generate-report")
async def generate_report(data: ReportData):
    # Choose the prompt language based on selection
    if data.language.lower() == "hindi":
        prompt = f"""
आप एक पेशेवर सहायक हैं जिसे औपचारिक उत्पीड़न रिपोर्ट तैयार करने का काम सौंपा गया है। एक संरचित रिपोर्ट लिखने के लिए निम्नलिखित इनपुट का उपयोग करें:
घटना की तारीख: {data.date}
घटना का समय: {data.time}
घटना का स्थान: {data.location}
शामिल व्यक्ति: {data.person}
रिपोर्टर का नाम: {data.reporter}
प्राप्तकर्ता: {data.recipient}
विवरण: {data.description}
साक्ष्य: {data.evidence}
निर्देश:
- विषय, परिचय, घटना विवरण, समर्थन साक्ष्य (यदि प्रदान किया गया हो), कार्रवाई के लिए अनुरोध, समापन कथन, हस्ताक्षर शामिल करें।
- पेशेवर टोन बनाए रखें।
रिपोर्ट पूरी तरह से हिंदी में लिखें।
-यदि उपयोगकर्ता सबूत (जैसे ईमेल, संदेश या सीसीटीवी फुटेज) प्रदान करता है, तो "समर्थक साक्ष्य" अनुभाग शामिल करें।

यदि उपयोगकर्ता कोई सबूत प्रदान नहीं करता है, तो "समर्थक साक्ष्य" अनुभाग को पूरी तरह हटा दें, उसके बारे में कोई पैराग्राफ भी न दें।

रिपोर्ट सीधे निर्दिष्ट प्राप्तकर्ता को संबोधित करें: {data.reporter}

नीचे दिए गए पत्र के उदाहरण के समान ही जनरेट करें:

विषय: औपचारिक उत्पीड़न शिकायत – 2025-05-08 की घटना

प्रति: [एचआर मैनेजर का नाम या मानव संसाधन विभाग]
प्रेषक: [आपका पूरा नाम]
कर्मचारी आईडी: [आपकी आईडी, यदि लागू हो]
दिनांक: 2025-05-08

परिचय:
मैं 8 मई 2025 को हुई कार्यस्थल पर उत्पीड़न की एक घटना की औपचारिक रूप से रिपोर्ट कर रहा/रही हूँ, जिसमें मेरे पर्यवेक्षक श्री राजेश शामिल थे। यह घटना माइक्रोसॉफ्ट के हैदराबाद कार्यालय में हुई। इस रिपोर्ट का उद्देश्य पूरी जांच और कंपनी की उत्पीड़न नीति के अनुसार उचित कार्रवाई प्राप्त करना है।

घटना का विवरण:
8 मई 2025 को लगभग शाम 5:39 बजे, श्री राजेश ने मुझे नियमित कार्य समय के बाद रुकने को कहा। इस दौरान उन्होंने ऐसा व्यवहार किया जिससे मेरी व्यक्तिगत सीमाएं लांघी गईं और यह उत्पीड़न की श्रेणी में आता है। उनके व्यवहार से मुझे असहज, अपमानित और असुरक्षित महसूस हुआ। मैं दृढ़ता से मानता/मानती हूँ कि यह घटना माइक्रोसॉफ्ट के आचार संहिता और उत्पीड़न रोकथाम नीतियों का उल्लंघन है।

समर्थक साक्ष्य:
मेरे पास [साक्ष्य का संक्षिप्त विवरण, जैसे ईमेल, संदेश, गवाही या सीसीटीवी फुटेज] है। घटना की संवेदनशील प्रकृति के कारण कुछ साक्ष्य सीधे साझा करना कठिन हो सकता है, लेकिन मैं किसी भी आधिकारिक जांच में पूरा सहयोग देने और आवश्यकतानुसार अधिक विवरण देने के लिए तैयार हूँ।

कार्रवाई का अनुरोध:
मैं निवेदन करता/करती हूँ कि इस मामले को गंभीरता से लिया जाए और शीघ्र जांच की जाए। मुझे माइक्रोसॉफ्ट की एक सुरक्षित, सम्मानजनक और समावेशी कार्यस्थल की प्रतिबद्धता पर विश्वास है और मैं अनुरोध करता/करती हूँ कि इस व्यवहार पर अनुशासनात्मक कार्रवाई की जाए ताकि भविष्य में ऐसी घटनाओं की पुनरावृत्ति न हो।

समापन वक्तव्य:
इस गंभीर मामले पर ध्यान देने के लिए धन्यवाद। मैं एक पेशेवर और सम्मानजनक कार्यस्थल बनाए रखने के लिए प्रतिबद्ध हूँ और आशा करता/करती हूँ कि हम मिलकर इन मूल्यों को कायम रख सकते हैं। यदि किसी और जानकारी या स्पष्टीकरण की आवश्यकता हो, तो कृपया मुझे बताएं।

सादर,
[आपका पूरा नाम]
[आपकी संपर्क जानकारी]
"""
    elif data.language.lower() == "telugu":
        prompt = f"""
మీరు అధికారిక వేధింపుల నివేదికను రూపొందించే బాధ్యతను కలిగి ఉన్న వృత్తిపరమైన సహాయకులు. క్రింది ఇన్‌పుట్‌లను ఉపయోగించి నిర్మాణాత్మక నివేదికను వ్రాయండి:
సంఘటన తేదీ: {data.date}
సంఘటన సమయం: {data.time}
సంఘటన స్థలం: {data.location}
పాల్గొన్న వ్యక్తి: {data.person}
నివేదికదారుని పేరు: {data.reporter}
గ్రహీత: {data.recipient}
వివరణ: {data.description}
సాక్ష్యం: {data.evidence}
సూచనలు:
- విషయం, పరిచయం, సంఘటన వివరాలు, మద్దతు సాక్ష్యం (అందించినట్లయితే), చర్య కోసం అభ్యర్థన, ముగింపు ప్రకటన, సంతకం చేర్చండి.
- వృత్తిపరమైన స్వరాన్ని నిర్వహించండి.
నివేదికను పూర్తిగా తెలుగులో వ్రాయండి.
-వాడుకదారు ఆధారాలు (ఇమెయిల్లు, సందేశాలు లేదా సీసీటీవీ ఫుటేజీ) అందించినట్లయితే, "మద్దతు ఆధారాలు" విభాగాన్ని చేర్చండి.

వాడుకదారు ఆధారాలు ఇవ్వకపోతే, "మద్దతు ఆధారాలు" విభాగాన్ని పూర్తిగా తొలగించండి. దాని గురించి పేరాగ్రాఫ్ కూడా రాయవద్దు.

రిపోర్టును నేరుగా పేర్కొన్న గ్రహీతకు (receiver): {data.reporter} పంపండి.

ఇదిగో ఈ క్రింది లేఖ ఒక ఉదాహరణ. ఇదే విధంగా రూపొందించండి:

విషయం: అధికారిక వేధింపుల ఫిర్యాదు 2025-05-08 న జరిగిన సంఘటన

కు: [హెచ్ఆర్ మేనేజర్ పేరు లేదా మానవ వనరుల విభాగం]
నుండి: [మీ పూర్తి పేరు]
ఉద్యోగి ఐడి: [మీ ఐడి ఉంటే]
తేదీ: 2025-05-08

పరిచయం:
2025 మే 8న మైక్రోసాఫ్ట్ హైదరాబాద్ కార్యాలయంలో నా మేనేజర్ అయిన శ్రీ రాజేష్ మిమ్మల్ని వేధించిన సంఘటన గురించి అధికారికంగా నివేదిస్తున్నాను. కంపెనీ యొక్క వేధింపుల నిబంధనలకు అనుగుణంగా సమగ్ర దర్యాప్తు మరియు తగిన చర్య తీసుకోవాలనే ఉద్దేశంతో ఈ నివేదికను అందిస్తున్నాను.

సంఘటన వివరాలు:
2025 మే 8 సాయంత్రం సుమారు 5:39 గంటల సమయంలో, శ్రీ రాజేష్ నన్ను సాధారణ పని గంటల తర్వాత ఆగమని అడిగారు. ఆ సమయంలో ఆయన నా వ్యక్తిగత పరిమితులను ఉల్లంఘించేలా వ్యవహరించారు, ఇది నా దృష్టిలో వేధింపుగా భావించదగినది. ఈ ప్రవర్తన వలన నాకు అసౌకర్యం, అపమానం, భద్రతలేమిగా అనిపించింది. ఇది మైక్రోసాఫ్ట్ ప్రవర్తనా నియమావళి మరియు వేధింపుల నివారణ విధానాలను ఉల్లంఘించిందని నమ్ముతున్నాను.

మద్దతు ఆధారాలు:
నా వద్ద [ఇమెయిల్లు, సందేశాలు, సాక్షుల వాంగ్మూలం లేదా సీసీటీవీ ఫుటేజీ వంటి] ఆధారాలు ఉన్నాయి. ఈ సంఘటన యొక్క సున్నిత స్వభావాన్ని పరిగణనలోకి తీసుకుంటే కొన్ని ఆధారాలను నేరుగా పంచుకోవడం కష్టంగా ఉండవచ్చు, కానీ అధికార దర్యాప్తులో పూర్తి సహకారం అందించేందుకు నేను సిద్ధంగా ఉన్నాను.

చర్యకు అభ్యర్థన:
ఈ అంశాన్ని గంభీరంగా తీసుకుని వెంటనే దర్యాప్తు చేయమని నేను వినమ్రంగా అభ్యర్థిస్తున్నాను. మైక్రోసాఫ్ట్ ఒక సురక్షితమైన, గౌరవప్రదమైన మరియు సమావేశాత్మకమైన పని వాతావరణాన్ని కల్పించాలనే తన నిబద్ధతపై నాకు నమ్మకం ఉంది. ఈ ప్రవర్తనపై తగిన చర్యలు తీసుకోవాలని, భవిష్యత్తులో ఇటువంటి సంఘటనలు జరగకుండా చూడాలని కోరుతున్నాను.

ముగింపు ప్రకటన:
ఈ తీవ్రమైన విషయంపై మీరు చూపుతున్న శ్రద్ధకు ధన్యవాదాలు. ఒక వృత్తిపరమైన, గౌరవప్రదమైన పని వాతావరణాన్ని కొనసాగించాలన్న నా నిబద్ధతను తెలియజేస్తున్నాను. మరింత సమాచారం లేదా స్పష్టత అవసరమైతే, దయచేసి నాకు తెలియజేయండి.

భవదీయుడు/భవదీయురాలు,
[మీ పూర్తి పేరు]
[మీ సంప్రదింపు వివరాలు]
"""
    elif data.language.lower() == "english":  # Default to English
        prompt = f"""
You are a professional assistant tasked with generating a formal harassment report. Use the following inputs to write a structured report:
Date of Incident: {data.date}
Time of Incident: {data.time}
Place of Incident: {data.location}
Person Involved: {data.person}
Name of Reporter: {data.reporter}
Recipient: {data.recipient}
Description: {data.description}
Evidence: {data.evidence}
Instructions:
- Include a Subject, Introduction,use Respected as salutation Incident Details, Supporting Evidence (if provided), Request for Action, Closing Statement, Signature.
- Maintain professional tone.
-Note:

If the user provides evidence (such as emails, messages, or CCTV footage), include the "Supporting Evidence" section.

If the user does not provide evidence, omit the "Supporting Evidence" section entirely, do not even give a paragraph about it.

Address the report directly to the specified recipient: {data.reporter}

Below is an example of the letter, generate in similar way

Subject: Formal Harassment Complaint – Incident on 2025-05-08

To: [HR Manager’s Name or Human Resources Department]
From: [Your Full Name]
Employee ID: [Your ID, if applicable]
Date: 2025-05-08

Introduction:
I am writing to formally report an incident of workplace harassment that occurred on May 8, 2025, involving my supervisor, Mr. Rajesh, at the Microsoft office in Hyderabad. This report is submitted with the intent to seek a thorough investigation and appropriate action in accordance with the company’s harassment policy.

Details of the Incident:
On May 8, 2025, at approximately 5:39 PM, I was asked by Mr. Rajesh to stay back after regular working hours. During this time, Mr. Rajesh engaged in conduct that I believe violated my personal boundaries and constituted harassment. His behavior made me feel uncomfortable, disrespected, and unsafe in the workplace environment. I firmly believe this incident breaches Microsoft's code of conduct and harassment prevention policies.

Supporting Evidence:
I have [briefly describe the type of evidence, such as email correspondence, messages, witness testimony, or CCTV footage]. Due to the sensitive nature of the incident, some evidence may be challenging to share directly, but I am fully prepared to cooperate with any official investigation and provide additional details as required.

Request for Action:
I respectfully request that this matter be taken seriously and investigated promptly. I trust Microsoft’s commitment to fostering a safe, respectful, and inclusive workplace, and I urge that appropriate disciplinary actions be taken to address this behavior and to prevent similar occurrences in the future.

Closing Statement:
Thank you for your attention to this serious matter. I am committed to maintaining a professional and respectful workplace and hope we can work together to uphold these values. Please let me know if any further information or clarification is required.

Sincerely,
[Your Full Name]
[Your Contact Information]
"""
    
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": MODEL_NAME,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3
    }
    response = requests.post("https://api.groq.com/openai/v1/chat/completions", headers=headers, json=payload)
    if response.status_code == 200:
        report_text = response.json()['choices'][0]['message']['content']
        return {"report": report_text.strip()}
    return JSONResponse(status_code=500, content={"error": "Groq API call failed"})

@app.post("/download-report")
async def download_report(req: Request):
    body = await req.json()
    report = body.get("report", "")
    doc = Document()
    doc.add_heading("Harassment Report", level=1)
    doc.add_paragraph(report)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(buf, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers={"Content-Disposition": "attachment; filename=harassment_report.docx"})

@app.post("/send-email")
async def send_email(data: EmailData):
    try:
        email_list = [email.strip() for email in data.emails.split(',') if email.strip()]
        doc = Document()
        doc.add_heading("Harassment Report", level=1)
        doc.add_paragraph(data.report)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        for email in email_list:
            msg = EmailMessage()
            msg['Subject'] = "Harassment Report Submission"
            msg['From'] = SENDER_EMAIL
            msg['To'] = email
            msg.set_content("Please find the attached harassment report.")
            msg.add_attachment(buf.getvalue(), maintype='application', subtype='vnd.openxmlformats-officedocument.wordprocessingml.document', filename="harassment_report.docx")

            with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
                server.login(SENDER_EMAIL, SENDER_APP_PASSWORD)
                server.send_message(msg)

        return {"message": "Email(s) sent successfully."}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.post("/send-email-with-attachments")
async def send_email_with_attachments(
    report: str = Form(...),
    emails: str = Form(...),
    files: List[UploadFile] = File(None),
    photos: List[UploadFile] = File(None),
    audio: List[UploadFile] = File(None)
):
    try:
        # Process email recipients
        email_list = [email.strip() for email in emails.split(',') if email.strip()]
        
        # Create Word document from report
        doc = Document()
        doc.add_heading("Harassment Report", level=1)
        doc.add_paragraph(report)
        doc_buf = BytesIO()
        doc.save(doc_buf)
        doc_buf.seek(0)
        
        for email in email_list:
            # Create email message
            msg = EmailMessage()
            msg['Subject'] = "Harassment Report Submission"
            msg['From'] = SENDER_EMAIL
            msg['To'] = email
            
            # Add email body text
            msg.set_content("""
            Please find the attached harassment report and supporting files.
            
            This is an automated message sent from the H.E.A.R reporting system.
            """)
            
            # Add Word document attachment
            msg.add_attachment(doc_buf.getvalue(), 
                              maintype='application', 
                              subtype='vnd.openxmlformats-officedocument.wordprocessingml.document', 
                              filename="harassment_report.docx")
            
            # Add document file attachments
            if files:
                for file in files:
                    file_content = await file.read()
                    content_type = file.content_type or 'application/octet-stream'
                    maintype, subtype = content_type.split('/', 1)
                    msg.add_attachment(file_content, 
                                      maintype=maintype, 
                                      subtype=subtype, 
                                      filename=file.filename)
            
            # Add photo attachments
            if photos:
                for photo in photos:
                    photo_content = await photo.read()
                    content_type = photo.content_type or 'image/jpeg'
                    maintype, subtype = content_type.split('/', 1)
                    msg.add_attachment(photo_content, 
                                      maintype=maintype, 
                                      subtype=subtype, 
                                      filename=photo.filename)
            
            # Add audio attachments
            if audio:
                for audio_file in audio:
                    audio_content = await audio_file.read()
                    content_type = audio_file.content_type or 'audio/mpeg'
                    maintype, subtype = content_type.split('/', 1)
                    msg.add_attachment(audio_content, 
                                      maintype=maintype, 
                                      subtype=subtype, 
                                      filename=audio_file.filename)
            
            # Send the email
            with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
                server.login(SENDER_EMAIL, SENDER_APP_PASSWORD)
                server.send_message(msg)
                
        return {"message": f"Email(s) sent successfully to {len(email_list)} recipient(s) with attachments."}
    
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})